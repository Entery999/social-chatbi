from langchain.tools import tool
from app.utils.logger import Logger
from docx import Document
import time
from app.ai.schema.docxwriteResponse import Args
from pathlib import Path

logger = Logger().get_Logger(__name__)

@tool("docx_write_tool", args_schema=Args)
def docx_write_tool(content: str) -> str:
    """
    写入docx文档，保存到static/download/，返回下载路径
    """
    doc = Document()
    doc.add_heading("分析报告", level=1)
    doc.add_paragraph(content)
    file_name = time.strftime("%Y%m%d_%H%M%S")
    BASE_DIR = Path(__file__).resolve().parent.parent.parent
    log_dir = BASE_DIR / "static/download"
    log_dir.mkdir(parents=True, exist_ok=True)
    file_path = log_dir / f"{file_name}.docx"
    doc.save(file_path)
    down_path = f"http://localhost:8000/static/download/{file_name}.docx"
    logger.info(f"[docx_write] 已写入 {file_path}")
    return f"下载地址：{down_path}"
