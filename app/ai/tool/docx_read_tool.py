from langchain.tools import tool
from app.utils.logger import Logger
from docx import Document

from app.ai.schema.docxreadResponse import Args
logger = Logger().get_Logger(__name__)

@tool("docx_read_tool", args_schema=Args)
def docx_read_tool(path: str) -> str:
    """
    读取docx文件，返回所有段落文本和表格数据
    """
    doc = Document(path)
    lines = []
    # 读取段落
    for para in doc.paragraphs:
        if para.text.strip():
            lines.append(para.text.strip())
    # 读取表格
    for i, table in enumerate(doc.tables):
        lines.append(f"--- 表格{i+1} ---")
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            lines.append(" | ".join(cells))
    result = "\n".join(lines)
    logger.info(f"[docx_read] 读取{path}完成，共{len(lines)}行")
    return result



